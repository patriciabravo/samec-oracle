import os
import secrets
from . import acreditacion_bp
from flask import render_template, jsonify, request
from flask_login import login_required, current_user
from app import db
from flask import send_file
from docx import Document
from docxtpl import DocxTemplate
from io import BytesIO
from app.models import AcreditacionH1EquipoAcreditacion, Autoevaluacion, IpressEssalud, RedEssalud
from app.blueprints.acreditacion.service import AcreditacionHitoService
from app.constants import ACREDITACION_ACTUAL

@acreditacion_bp.route("/listado/hitos", methods=["GET"])
@login_required
def listado_hitos():
    id_ipress = request.form.get("id_ipress")
    return render_template(
        "listahitos.html",
        user=current_user,
        page_title="Reporte de Hitos",
        id_autoevaluacion=current_user.id_autoevaluacion,
        id_ipress=current_user.id_ipress
    )   
 
@acreditacion_bp.route("/api/hito-1/guardar",methods=["POST"])
@login_required
def guardar_hito_1():
    try:
        id_autoevaluacion = request.form.get("id_autoevaluacion")
        if not id_autoevaluacion:
            return jsonify({
                "success": False,
                "message": "Falta id_autoevaluacion"
            }), 400

        id_acreditacion = id_autoevaluacion
        AcreditacionHitoService.guardar_hito_1(
            id_acreditacion,
            request.form
        )
        return jsonify({
            "success": True,
            "message": "Hito 1 guardado correctamente"
        })
    except Exception as e:
        return jsonify({
            "success": False,
            "message": str(e)
        }), 500

@acreditacion_bp.route('/hito1/<int:id_acreditacion>',methods=['GET'])
@login_required
def editar_hito_1(id_acreditacion):
    response = AcreditacionHitoService.obtener_hito_1(id_acreditacion)
    return jsonify({
        "success": True,
        "data": response
    }), 200

@acreditacion_bp.route("/api/hito-1/word/<int:id_autoevaluacion>")
@login_required
def generar_word_hito_1(id_autoevaluacion):
    try:
        registro = (
            db.session.query(
                Autoevaluacion,
                IpressEssalud.nombre_ipress,
                RedEssalud.nombre_red
            )
            .join(IpressEssalud, Autoevaluacion.id_ipress == IpressEssalud.id_ipress)
            .join(RedEssalud, IpressEssalud.id_red == RedEssalud.id_red)           
            .filter(Autoevaluacion.id_autoevaluacion == id_autoevaluacion)
            .first()
        )

        template_path = os.path.join(os.path.dirname(__file__),"..","..","documents","templates","template_hito1.docx")
        doc = Document(template_path)

        for p in doc.paragraphs:
            if "{{red_prestacional}}" in p.text:
                p.text = p.text.replace(
                    "{{red_prestacional}}",
                    registro[2]
                )
            if "{{ipress_name}}" in p.text:
                p.text = p.text.replace(
                    "{{ipress_name}}",
                    registro[1]
                )
            if "{{anio_acreditacion}}" in p.text:
                p.text = p.text.replace(
                    "{{anio_acreditacion}}",
                    str(ACREDITACION_ACTUAL)
                )

        table = doc.tables[0]  # asumimos que es la primera tabla
        for i in range(len(table.rows) - 1, 0, -1):
            table._tbl.remove(table.rows[i]._tr)

        registros = AcreditacionH1EquipoAcreditacion.query.filter_by(id_acreditacion=id_autoevaluacion).all()        
        for r in registros:
            row = table.add_row().cells
            row[0].text = r.nombre_miembro
            row[1].text = r.cargo_miembro
            row[2].text = "SI" if r.es_lider else "NO"

        file_stream = BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        return send_file(file_stream,as_attachment=True,download_name="hito_1_generado.docx",mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    except Exception as e:
        return {
            "success": False,
            "message": str(e)
        }, 500
