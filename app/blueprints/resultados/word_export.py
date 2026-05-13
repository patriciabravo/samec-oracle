from collections import OrderedDict
from datetime import date
from io import BytesIO

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT


def _aplicar_orientacion_horizontal(doc):
    for section in doc.sections:
        w, h = section.page_width, section.page_height
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = h
        section.page_height = w


def _puntaje_a_texto(puntaje):
    try:
        p = int(puntaje)
    except (TypeError, ValueError):
        return "No cumple"
    if p == 2:
        return "Cumple Totalmente"
    if p == 1:
        return "Cumple parcialmente"
    return "No cumple"


def _agrupar_por_estandar_y_criterio(filas):
    """
    Agrupa por estándar (DIR-1) y dentro arma filas por criterio (DIR-1-1).
    Cada fila del API puede venir repetida por múltiples fuentes/técnicas: se agrupa por código de criterio.
    """
    orden_estandares = []
    por_estandar = OrderedDict()

    for row in filas:
        mp = row.get("nombre_mp") or "(Sin macroproceso)"
        est = row.get("codigo_estandar")
        cod_criterio = row.get("codigo_criterio")
        if est is None or not str(est).strip() or not cod_criterio:
            continue

        est = str(est).strip()
        if est not in por_estandar:
            orden_estandares.append(est)
            por_estandar[est] = {
                "macroproceso": mp,
                "criterios": OrderedDict(),
            }

        bloque = por_estandar[est]
        criterios = bloque["criterios"]

        if cod_criterio not in criterios:
            criterios[cod_criterio] = {
                "puntaje": row.get("puntaje_criterio"),
                "fuentes": set(),
                "tecnicas": set(),
            }

        celda = criterios[cod_criterio]
        if celda["puntaje"] is None and row.get("puntaje_criterio") is not None:
            celda["puntaje"] = row.get("puntaje_criterio")

        nf = row.get("nombre_fuente")
        if nf and str(nf).strip():
            celda["fuentes"].add(str(nf).strip())

        nt = row.get("nombre_tecnica")
        if nt and str(nt).strip():
            celda["tecnicas"].add(str(nt).strip())

    return [(est, por_estandar[est]) for est in orden_estandares]


def generar_word_autoevaluacion(filas, nombre_ipress=""):
    doc = Document()
    _aplicar_orientacion_horizontal(doc)
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    titulo = doc.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = titulo.add_run("Hoja de Registro de Datos para Autoevaluación")
    run.bold = True
    run.font.size = Pt(14)
    if nombre_ipress:
        sub = doc.add_paragraph()
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub.add_run(nombre_ipress).italic = True

    doc.add_paragraph()
    bloques = _agrupar_por_estandar_y_criterio(filas)
    first_est = True
    fecha_hoy = date.today().strftime("%d/%m/%Y")

    labels_encabezado = [
        "Establecimiento de Salud",
        "Código del Estándar",
        "Macroproceso",
        "Evaluador(es)",
        "Fecha",
        "Servicio Evaluados",
        "Participantes de la Evaluación",
    ]

    for codigo_estandar, bloque in bloques:
        criterios = bloque["criterios"]
        if not criterios:
            continue
        if not first_est:
            doc.add_page_break()
        first_est = False

        enc = doc.add_table(rows=len(labels_encabezado), cols=2)
        enc.style = "Table Grid"
        valores_enc = [
            "",
            codigo_estandar,
            bloque["macroproceso"] or "",
            "",
            fecha_hoy,
            "",
            "",
        ]
        for i, lab in enumerate(labels_encabezado):
            enc.rows[i].cells[0].text = lab + ":"
            enc.rows[i].cells[1].text = valores_enc[i]

        doc.add_paragraph()

        tbl = doc.add_table(rows=1 + len(criterios), cols=5)
        tbl.style = "Table Grid"
        hdr = tbl.rows[0].cells
        headers = [
            "Código de criterio de evaluación",
            "Puntaje de autoevaluación",
            "Fuente auditable",
            "Técnica utilizada",
            "Sustento del puntaje / Resultado",
        ]
        for j, h in enumerate(headers):
            hdr[j].text = h
            for p in hdr[j].paragraphs:
                for r in p.runs:
                    r.bold = True

        for idx, (codigo, info) in enumerate(bloque["criterios"].items(), start=1):
            row = tbl.rows[idx].cells
            pje = info["puntaje"]
            row[0].text = str(codigo) if codigo else ""
            if pje is None:
                row[1].text = ""
            else:
                try:
                    row[1].text = str(int(pje))
                except (TypeError, ValueError):
                    row[1].text = str(pje)
            fuentes = sorted(info["fuentes"])
            row[2].text = " / ".join(fuentes) if fuentes else ""
            tecs = sorted(info["tecnicas"])
            row[3].text = " / ".join(tecs) if tecs else ""
            row[4].text = _puntaje_a_texto(pje)

        doc.add_paragraph()

    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio
